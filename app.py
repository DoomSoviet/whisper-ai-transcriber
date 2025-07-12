#!/usr/bin/env python3
"""
Whisper AI Transcription Service
A lightweight web application for transcribing audio from YouTube videos and local files.
"""

import os
import tempfile
import shutil
from flask import Flask, request, render_template, jsonify, send_file
import whisper
import yt_dlp
from werkzeug.utils import secure_filename
import logging
from datetime import datetime
import json
from docx import Document
import threading

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['TRANSCRIPTS_FOLDER'] = 'transcripts'

# Create necessary directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['TRANSCRIPTS_FOLDER'], exist_ok=True)

# Supported file extensions
ALLOWED_EXTENSIONS = {'mp3', 'mp4', 'wav', 'flac', 'm4a', 'ogg', 'wma', 'aac'}

# Global variable to store the Whisper model
whisper_model = None

# Global variable to track the current transcription thread
current_transcription_thread = None
cancel_event = threading.Event()

def load_whisper_model(model_size="base"):
    """Load the Whisper model."""
    global whisper_model
    if whisper_model is None:
        logger.info(f"Loading Whisper model: {model_size}")
        try:
            whisper_model = whisper.load_model(model_size)
            logger.info("Whisper model loaded successfully")
        except Exception as e:
            logger.error(f"Error loading Whisper model: {e}")
            raise
    return whisper_model

def allowed_file(filename):
    """Check if the uploaded file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def download_youtube_audio(url, output_path):
    """Download audio from YouTube video."""
    try:
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': os.path.join(output_path, '%(title)s.%(ext)s'),
            'extractaudio': True,
            'audioformat': 'mp3',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            # Add user agent and other headers to avoid 403 errors
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            },
            # Additional options to handle restrictions
            'extractor_retries': 3,
            'fragment_retries': 3,
            'skip_unavailable_fragments': True,
            'ignore_errors': False,
            'no_warnings': False,
            # Use cookies if available
            'cookiefile': None,
            # Throttle requests to avoid rate limiting
            'sleep_interval': 1,
            'max_sleep_interval': 5,
            # Try different extraction methods
            'youtube_include_dash_manifest': False,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            try:
                # First try to get info without downloading
                info = ydl.extract_info(url, download=False)
                title = info.get('title', 'Unknown')
                
                # Then download
                ydl.download([url])
                
                # Find the downloaded file
                for file in os.listdir(output_path):
                    if file.endswith('.mp3'):
                        return os.path.join(output_path, file), title
                        
            except yt_dlp.utils.DownloadError as e:
                if "403" in str(e) or "Forbidden" in str(e):
                    raise Exception("YouTube access forbidden. This video may be restricted or require sign-in. Try a different video or check if the URL is correct.")
                elif "unavailable" in str(e).lower():
                    raise Exception("This YouTube video is unavailable or private. Please try a different video.")
                elif "ffmpeg" in str(e).lower() or "ffprobe" in str(e).lower():
                    raise Exception("FFmpeg is required for YouTube downloads but not found. Please restart your terminal/browser and try again, or install FFmpeg manually.")
                else:
                    raise Exception(f"YouTube download failed: {str(e)}")
                    
        return None, None
    except Exception as e:
        logger.error(f"Error downloading YouTube video: {e}")
        raise

def transcribe_audio(audio_path, model_size="base"):
    """Transcribe audio file using Whisper, with cancellation support."""
    try:
        model = load_whisper_model(model_size)
        logger.info(f"Starting transcription of: {audio_path}")
        # Check for cancellation before starting
        if cancel_event.is_set():
            raise Exception("Transcription cancelled by user.")
        result = model.transcribe(audio_path)
        # Check for cancellation after transcription (if possible)
        if cancel_event.is_set():
            raise Exception("Transcription cancelled by user.")
        return {
            'text': result['text'],
            'segments': result['segments'],
            'language': result['language']
        }
    except Exception as e:
        logger.error(f"Error during transcription: {e}")
        raise
    finally:
        cancel_event.clear()

def save_transcript(transcript_data, filename_base):
    """Save transcript to file."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    transcript_filename = f"{filename_base}_{timestamp}.json"
    transcript_path = os.path.join(app.config['TRANSCRIPTS_FOLDER'], transcript_filename)
    
    with open(transcript_path, 'w', encoding='utf-8') as f:
        json.dump(transcript_data, f, ensure_ascii=False, indent=2)
    
    return transcript_path

@app.route('/')
def index():
    """Main page."""
    return render_template('index.html')

@app.route('/transcribe_youtube', methods=['POST'])
def transcribe_youtube():
    """Transcribe audio from YouTube video."""
    try:
        data = request.get_json()
        youtube_url = data.get('url')
        model_size = data.get('model_size', 'base')
        
        if not youtube_url:
            return jsonify({'error': 'YouTube URL is required'}), 400
        
        # Create temporary directory for download
        with tempfile.TemporaryDirectory() as temp_dir:
            # Download audio from YouTube
            audio_path, video_title = download_youtube_audio(youtube_url, temp_dir)
            
            if not audio_path:
                return jsonify({'error': 'Failed to download audio from YouTube'}), 400
            
            # Transcribe the audio
            transcript = transcribe_audio(audio_path, model_size)
            
            # Add metadata
            transcript['source'] = 'youtube'
            transcript['url'] = youtube_url
            transcript['title'] = video_title
            transcript['timestamp'] = datetime.now().isoformat()
            
            # Save transcript
            safe_title = secure_filename(video_title or 'youtube_video')
            transcript_path = save_transcript(transcript, safe_title)
            
            return jsonify({
                'success': True,
                'transcript': transcript['text'],
                'language': transcript['language'],
                'title': video_title,
                'transcript_file': os.path.basename(transcript_path)
            })
            
    except Exception as e:
        logger.error(f"Error in YouTube transcription: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/transcribe_file', methods=['POST'])
def transcribe_file():
    """Transcribe uploaded audio file."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        model_size = request.form.get('model_size', 'base')
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not supported'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Transcribe the audio
            transcript = transcribe_audio(file_path, model_size)
            
            # Add metadata
            transcript['source'] = 'file'
            transcript['filename'] = filename
            transcript['timestamp'] = datetime.now().isoformat()
            
            # Save transcript
            filename_base = os.path.splitext(filename)[0]
            transcript_path = save_transcript(transcript, filename_base)
            
            return jsonify({
                'success': True,
                'transcript': transcript['text'],
                'language': transcript['language'],
                'filename': filename,
                'transcript_file': os.path.basename(transcript_path)
            })
            
        finally:
            # Clean up uploaded file
            if os.path.exists(file_path):
                os.remove(file_path)
                
    except Exception as e:
        logger.error(f"Error in file transcription: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_transcript/<filename>')
def download_transcript(filename):
    """Download saved transcript file as JSON, TXT, or DOCX."""
    try:
        transcript_path = os.path.join(app.config['TRANSCRIPTS_FOLDER'], filename)
        if not os.path.exists(transcript_path):
            return jsonify({'error': 'Transcript file not found'}), 404

        # Determine export format from query param
        export_format = request.args.get('format', 'json').lower()
        base_name = os.path.splitext(filename)[0]

        # Load transcript data
        with open(transcript_path, 'r', encoding='utf-8') as f:
            transcript_data = json.load(f)

        if export_format == 'txt':
            txt_path = os.path.join(app.config['TRANSCRIPTS_FOLDER'], base_name + '.txt')
            with open(txt_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(transcript_data['text'])
            return send_file(txt_path, as_attachment=True)

        elif export_format == 'docx':
            docx_path = os.path.join(app.config['TRANSCRIPTS_FOLDER'], base_name + '.docx')
            doc = Document()
            doc.add_heading('Transcript', 0)
            doc.add_paragraph(transcript_data['text'])
            doc.save(docx_path)
            return send_file(docx_path, as_attachment=True)

        else:  # Default: JSON
            return send_file(transcript_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/health')
def health_check():
    """Health check endpoint."""
    try:
        # Try to load the model to verify everything is working
        load_whisper_model()
        return jsonify({'status': 'healthy', 'whisper_loaded': True})
    except Exception as e:
        return jsonify({'status': 'unhealthy', 'error': str(e)}), 500

@app.route('/cancel_transcription', methods=['POST'])
def cancel_transcription():
    """Cancel the current transcription process safely."""
    global cancel_event
    cancel_event.set()
    return jsonify({'success': True, 'message': 'Transcription cancelled.'})

if __name__ == '__main__':
    print("Starting Whisper Transcription Service...")
    print("Loading Whisper model (this may take a moment)...")
    
    try:
        load_whisper_model()
        print("✓ Whisper model loaded successfully")
        print("✓ Server starting on http://localhost:5000")
        app.run(debug=True, host='0.0.0.0', port=5000)
    except Exception as e:
        print(f"✗ Error starting server: {e}")
        print("Please make sure all dependencies are installed (run: pip install -r requirements.txt)")
